import os
import re
import base64
import json
import io
from flask import Flask, render_template, request, redirect, url_for, send_file, jsonify, session
from flask_cors import CORS
from werkzeug.utils import secure_filename
import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import fitz
import tempfile

# --- CẤU HÌNH FLASK ---
app = Flask(__name__)
app.secret_key = '1ae05550c898138fc632e4e6c0fba3f14cc10104e5697f19eb6fde9467b8d0cd19ab1faaa659f982a4c479d7f3d8827f815043d5064bec6b0c1d6e45842b77a'  # Thay đổi thành secret key thực tế

# Cấu hình CORS cho Next.js
try:
    from flask_cors import CORS
    CORS(app, origins=["http://localhost:3000"])  # Next.js development server
except ImportError:
    print("Warning: flask-cors not installed. CORS may not work properly.")

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# --- HÀM HELPER ---
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# --- LOGIC XỬ LÝ DOCX (PHIÊN BẢN CUỐI CÙNG, HOÀN CHỈNH) ---
def extract_docx_data(file_path):
    """
    Hàm xử lý DOCX cuối cùng:
    1. Phân tích cấu trúc câu hỏi/đáp án (đã hoạt động tốt).
    2. Thêm một bước riêng để tìm highlight và cập nhật đáp án đúng.
    """
    try:
        doc = docx.Document(file_path)
        quiz_data = []

        # --- Nhiệm vụ 1: Lấy cấu trúc câu hỏi và các lựa chọn ---
        # Logic này được giữ lại từ phiên bản đã hoạt động tốt của bạn
        
        # Chiến lược 1: Xử lý file có định dạng mỗi câu hỏi/đáp án trên 1 dòng
        current_quiz = None
        for p in doc.paragraphs:
            p_text = p.text.strip()
            if not p_text: continue

            match_question = re.match(r'^(Câu\s*(\d+)[:.]?)(.*)', p_text, re.IGNORECASE)
            if match_question:
                if current_quiz: quiz_data.append(current_quiz)
                current_quiz = {"question_number": int(match_question.group(2)), "question_text": match_question.group(0), "options": [], "correct_answer_index": -1, "correct_answer_char": '', "images": []}
            elif current_quiz:
                if re.search(r'^[A-D][.:]', p_text):
                    current_quiz["options"].append(p_text)
                else:
                    current_quiz["question_text"] += "\n" + p_text
        if current_quiz: quiz_data.append(current_quiz)

        # Kiểm tra xem chiến lược 1 có hiệu quả không
        # Nếu không có lựa chọn nào được tìm thấy, thử chiến lược 2
        if not any(q['options'] for q in quiz_data):
            print("Fallback to single-line format strategy")
            quiz_data = [] # Reset và bắt đầu lại
            for p in doc.paragraphs:
                p_text = p.text.strip()
                if not p_text or not p_text.startswith("Câu"): continue
                
                question_blocks = re.split(r'(Câu\s*(\d+)[:.]?)', p_text, flags=re.IGNORECASE)
                if question_blocks[0] == '': question_blocks = question_blocks[1:]
                
                for i in range(0, len(question_blocks), 3):
                    if i + 2 >= len(question_blocks): continue
                    
                    question_number = int(question_blocks[i+1])
                    full_question_text = (question_blocks[i] + question_blocks[i+2]).strip()
                    parts = re.split(r'(\s+[A-D][.:])', full_question_text, maxsplit=1, flags=re.IGNORECASE)
                    
                    question_text = parts[0].strip()
                    options_string = (parts[1] + parts[2]).strip() if len(parts) > 2 else ""
                    
                    pattern = re.compile(r'([A-D][.:])(.*?)(?=\s*[A-D][.:]|\s*Câu\s*\d+|\Z)', re.IGNORECASE | re.DOTALL)
                    matches = pattern.findall(options_string)
                    options = [f"{m[0].strip()} {m[1].strip()}" for m in matches]
                    
                    quiz_data.append({"question_number": question_number, "question_text": question_text, "options": options, "correct_answer_index": -1, "correct_answer_char": '', "images": []})

        # --- Nhiệm vụ 2: Duyệt lại để tìm highlight và cập nhật đáp án đúng ---
        # Duyệt qua từng câu hỏi đã có
        for quiz in quiz_data:
            # Tìm paragraph tương ứng với câu hỏi này
            for p in doc.paragraphs:
                # Nếu text của paragraph chứa số câu hỏi và một phần nội dung
                if f"Câu {quiz['question_number']}" in p.text:
                    # Duyệt qua các run trong paragraph đó để tìm highlight
                    for run in p.runs:
                        if run.font.highlight_color is not None:
                            match_char = re.search(r'[A-D]', run.text, re.IGNORECASE)
                            if match_char:
                                quiz["correct_answer_char"] = match_char.group(0).upper()
                                break # Đã tìm thấy, thoát vòng lặp run
                    # Nếu đã tìm thấy đáp án cho câu này, thoát vòng lặp paragraph
                    if quiz["correct_answer_char"]:
                        break
        
        # --- Nhiệm vụ 3: Dọn dẹp cuối cùng và gán index ---
        for quiz in quiz_data:
            if quiz["options"]:
                try:
                    first_option_start_index = quiz["question_text"].find(quiz["options"][0])
                    if first_option_start_index != -1:
                        quiz["question_text"] = quiz["question_text"][:first_option_start_index].strip()
                except (IndexError, TypeError): pass
            
            if quiz["correct_answer_char"]:
                for i, opt in enumerate(quiz["options"]):
                    if opt.strip().upper().startswith(quiz["correct_answer_char"]):
                        quiz["correct_answer_index"] = i
                        break
        
        return quiz_data

    except Exception as e:
        print(f"Error reading docx: {str(e)}")
        return []

# --- Các hàm và route còn lại giữ nguyên ---
def parse_horizontal_options(text):
    pattern = re.compile(r'([A-D][.:])(.*?)(?=\s*[A-D][.:]|\s*Câu\s*\d+|\Z)', re.IGNORECASE | re.DOTALL)
    matches = pattern.findall(text)
    return [f"{m[0].strip()} {m[1].strip()}" for m in matches]

def extract_pdf_data(file_path):
    try:
        doc = fitz.open(file_path)
        full_text = ""
        images = []
        for page in doc:
            full_text += page.get_text("text") + "\n"
            for img in page.get_images(full=True):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_base64 = base64.b64encode(image_bytes).decode('utf-8')
                images.append(f"data:image/{base_image['ext']};base64,{image_base64}")
        
        quiz_data = []
        question_blocks = re.split(r'(Câu\s*(\d+)[:.]?)', full_text, flags=re.IGNORECASE | re.DOTALL)
        if len(question_blocks) < 3: return []
        question_blocks = question_blocks[1:]
        
        for i in range(0, len(question_blocks), 3):
            if i + 2 >= len(question_blocks): continue
            
            question_header = question_blocks[i]
            question_number = int(question_blocks[i+1])
            question_body = question_blocks[i+2].strip()
            
            parts = re.split(r'(?=\s*[A-D][.:])', question_body, 1)
            question_part = parts[0]
            options_part = parts[1] if len(parts) > 1 else ""

            options = parse_horizontal_options(options_part)
            
            full_question_text = (question_header + " " + question_part).strip()
            quiz_data.append({
                "question_number": question_number, "question_text": full_question_text,
                "options": options, "correct_answer_index": -1, "correct_answer_char": '', "images": []
            })
        
        if images:
             quiz_data.append({
                 "question_number": "Hình ảnh", "question_text": "CÁC HÌNH ẢNH TRÍCH XUẤT TỪ FILE PDF:",
                 "options": [], "images": images, "correct_answer_index": -1, "correct_answer_char": ""
             })
        return quiz_data
    except Exception as e:
        print(f"Lỗi khi xử lý PDF: {e}")
        return []

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        if 'file' not in request.files: return redirect(request.url)
        file = request.files['file']
        if file.filename == '' or not allowed_file(file.filename): return redirect(request.url)
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        quiz_data = []
        file_ext = filename.rsplit('.', 1)[1].lower()
        if file_ext == 'docx':
            quiz_data = extract_docx_data(file_path)
        elif file_ext == 'pdf':
            quiz_data = extract_pdf_data(file_path)
        os.remove(file_path)
        
        # Lưu dữ liệu vào session thay vì truyền qua form
        session['quiz_data'] = quiz_data
        session['filename'] = filename
        
        # Hãy chắc chắn bạn có file 'results.html' trong thư mục templates
        return render_template('results.html', quiz_data=quiz_data, filename=filename)
    # Hãy chắc chắn bạn có file 'index.html' trong thư mục templates
    return render_template('index.html')

@app.route('/api/extract-quiz', methods=['POST'])
def extract_quiz_api():
    """API endpoint cho Next.js để trích xuất câu hỏi từ file"""
    try:
        # Kiểm tra file upload
        if 'file' not in request.files:
            return jsonify({'error': 'Không có file được upload'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Không có file được chọn'}), 400
            
        if not allowed_file(file.filename):
            return jsonify({'error': 'Định dạng file không được hỗ trợ. Chỉ chấp nhận .pdf và .docx'}), 400
        
        # Lưu file tạm thời
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        try:
            # Trích xuất dữ liệu
            quiz_data = []
            file_ext = filename.rsplit('.', 1)[1].lower()
            
            if file_ext == 'docx':
                quiz_data = extract_docx_data(file_path)
            elif file_ext == 'pdf':
                quiz_data = extract_pdf_data(file_path)
            
            # Xóa file tạm
            os.remove(file_path)
            
            # Trả về dữ liệu JSON
            return jsonify({
                'success': True,
                'filename': filename,
                'quiz_data': quiz_data,
                'total_questions': len([q for q in quiz_data if isinstance(q.get('question_number'), int)])
            })
            
        except Exception as e:
            # Xóa file nếu có lỗi
            if os.path.exists(file_path):
                os.remove(file_path)
            raise e
            
    except Exception as e:
        print(f"Error in extract_quiz_api: {str(e)}")
        return jsonify({'error': f'Lỗi khi xử lý file: {str(e)}'}), 500

@app.route('/download/<format>', methods=['POST'])
def download_quiz(format):
    """Route để tải về đề thi theo định dạng được chọn"""
    try:
        # Lấy dữ liệu từ session thay vì form
        quiz_data = session.get('quiz_data')
        filename = session.get('filename', 'quiz')
        
        if not quiz_data:
            return jsonify({'error': 'Không có dữ liệu quiz trong session'}), 400
        
        if format == 'pdf':
            return download_quiz_pdf(quiz_data, filename)
        elif format == 'docx':
            return download_quiz_docx(quiz_data, filename)
        else:
            return jsonify({'error': 'Định dạng không hỗ trợ'}), 400
            
    except Exception as e:
        print(f"Error in download_quiz: {str(e)}")
        return jsonify({'error': f'Lỗi khi tạo file: {str(e)}'}), 500

def download_quiz_docx(quiz_data, original_filename):
    """Tạo file Word chứa đề thi"""
    try:
        # Tạo document mới
        doc = docx.Document()
        
        # Thêm tiêu đề
        title = doc.add_heading(f'Đề thi từ file: {original_filename}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Thêm các câu hỏi
        for item in quiz_data:
            if isinstance(item.get('question_number'), int):
                # Thêm câu hỏi
                question_para = doc.add_paragraph()
                question_para.add_run(f"Câu {item['question_number']}: ").bold = True
                question_para.add_run(item['question_text'])
                
                # Thêm các lựa chọn
                for option in item.get('options', []):
                    option_para = doc.add_paragraph(option)
                    option_para.style = 'List Bullet'
                
                # Thêm khoảng trắng
                doc.add_paragraph()
        
        # Thêm bảng đáp án
        if any(item.get('correct_answer_char') for item in quiz_data):
            doc.add_heading('Bảng đáp án', level=1)
            
            # Tạo bảng với 8 cột (4 cặp Câu-Đáp án)
            table = doc.add_table(rows=1, cols=8)
            table.style = 'Table Grid'
            
            # Header
            header_cells = table.rows[0].cells
            headers = ['Câu', 'Đáp án', 'Câu', 'Đáp án', 'Câu', 'Đáp án', 'Câu', 'Đáp án']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            # Thêm dữ liệu theo hàng (4 câu mỗi hàng)
            quiz_with_answers = [item for item in quiz_data if isinstance(item.get('question_number'), int)]
            for i in range(0, len(quiz_with_answers), 4):
                row_cells = table.add_row().cells
                for j in range(4):
                    if i + j < len(quiz_with_answers):
                        item = quiz_with_answers[i + j]
                        row_cells[j*2].text = str(item['question_number'])
                        row_cells[j*2 + 1].text = item.get('correct_answer_char', '?')
        
        # Lưu vào memory
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # Tạo tên file
        base_name = original_filename.rsplit('.', 1)[0] if '.' in original_filename else original_filename
        download_filename = f"{base_name}_quiz.docx"
        
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=download_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        print(f"Error creating DOCX: {str(e)}")
        return jsonify({'error': f'Lỗi khi tạo file Word: {str(e)}'}), 500

def download_quiz_pdf(quiz_data, original_filename):
    """Tạo file PDF chứa đề thi với hỗ trợ tiếng Việt"""
    try:
        # Tạo file tạm thời
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            pdf_path = tmp_file.name
        
        # Kiểm tra xem reportlab có được cài đặt không
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            from reportlab.lib.units import inch
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
        except ImportError:
            return jsonify({'error': 'Thư viện reportlab chưa được cài đặt. Vui lòng cài đặt bằng: pip install reportlab'}), 500
        
        # Tạo PDF với encoding UTF-8
        c = canvas.Canvas(pdf_path, pagesize=A4)
        width, height = A4
        
        # Đăng ký font hỗ trợ tiếng Việt (Arial hoặc Times New Roman)
        font_name = 'Arial'
        font_bold = 'ArialBold'
        try:
            # Đăng ký font Arial từ hệ thống Windows
            system_font_paths = [
                'C:/Windows/Fonts/arial.ttf',      # Arial
                'C:/Windows/Fonts/arialbd.ttf',    # Arial Bold
                'C:/Windows/Fonts/times.ttf',      # Times New Roman
                'C:/Windows/Fonts/timesbd.ttf',    # Times New Roman Bold
            ]
            
            font_registered = False
            for sys_font in system_font_paths:
                if os.path.exists(sys_font):
                    try:
                        if 'arialbd' in sys_font.lower():
                            pdfmetrics.registerFont(TTFont('ArialBold', sys_font))
                            font_bold = 'ArialBold'
                        elif 'arial' in sys_font.lower():
                            pdfmetrics.registerFont(TTFont('Arial', sys_font))
                            font_name = 'Arial'
                            font_registered = True
                        elif 'timesbd' in sys_font.lower():
                            pdfmetrics.registerFont(TTFont('TimesBold', sys_font))
                            if not font_registered:
                                font_bold = 'TimesBold'
                        elif 'times' in sys_font.lower():
                            pdfmetrics.registerFont(TTFont('Times', sys_font))
                            if not font_registered:
                                font_name = 'Times'
                                font_registered = True
                    except Exception as e:
                        print(f"Error registering font {sys_font}: {e}")
                        continue
            
            # Nếu không đăng ký được font, dùng fallback font Helvetica
            if not font_registered:
                font_name = 'Helvetica'
                font_bold = 'Helvetica-Bold'
                print("Warning: Using fallback font Helvetica, Vietnamese characters may not display correctly")
                        
        except Exception as e:
            print(f"Font setup error: {e}")
            font_name = 'Helvetica'
            font_bold = 'Helvetica-Bold'
        
        def safe_text(text):
            """Hàm helper để xử lý text an toàn cho PDF"""
            if not text:
                return ""
            try:
                # Kiểm tra xem font có hỗ trợ ký tự Unicode không
                # Nếu dùng Arial hoặc Times, không cần thay thế ký tự
                if font_name in ['Arial', 'Times']:
                    return text
                # Fallback cho font không hỗ trợ Unicode
                vietnamese_map = {
                    'ế': 'e', 'ề': 'e', 'ể': 'e', 'ễ': 'e', 'ệ': 'e',
                    'ố': 'o', 'ồ': 'o', 'ổ': 'o', 'ỗ': 'o', 'ộ': 'o',
                    'ứ': 'u', 'ừ': 'u', 'ử': 'u', 'ữ': 'u', 'ự': 'u',
                    'ấ': 'a', 'ầ': 'a', 'ẩ': 'a', 'ẫ': 'a', 'ậ': 'a',
                    'ắ': 'a', 'ằ': 'a', 'ẳ': 'a', 'ẵ': 'a', 'ặ': 'a',
                    'ì': 'i', 'í': 'i', 'ỉ': 'i', 'ĩ': 'i', 'ị': 'i',
                    'ỳ': 'y', 'ý': 'y', 'ỷ': 'y', 'ỹ': 'y', 'ỵ': 'y',
                    'à': 'a', 'á': 'a', 'ả': 'a', 'ã': 'a', 'ạ': 'a',
                    'è': 'e', 'é': 'e', 'ẻ': 'e', 'ẽ': 'e', 'ẹ': 'e',
                    'ò': 'o', 'ó': 'o', 'ỏ': 'o', 'õ': 'o', 'ọ': 'o',
                    'ù': 'u', 'ú': 'u', 'ủ': 'u', 'ũ': 'u', 'ụ': 'u',
                    'ô': 'o', 'ê': 'e', 'â': 'a', 'ă': 'a', 'ư': 'u',
                    'đ': 'd', 'Đ': 'D',
                    'À': 'A', 'Á': 'A', 'Ả': 'A', 'Ã': 'A', 'Ạ': 'A',
                    'Ằ': 'A', 'Ắ': 'A', 'Ẳ': 'A', 'Ẵ': 'A', 'Ặ': 'A',
                    'Ầ': 'A', 'Ấ': 'A', 'Ẩ': 'A', 'Ẫ': 'A', 'Ậ': 'A',
                    'È': 'E', 'É': 'E', 'Ẻ': 'E', 'Ẽ': 'E', 'Ẹ': 'E',
                    'Ề': 'E', 'Ế': 'E', 'Ể': 'E', 'Ễ': 'E', 'Ệ': 'E',
                    'Ì': 'I', 'Í': 'I', 'Ỉ': 'I', 'Ĩ': 'I', 'Ị': 'I',
                    'Ò': 'O', 'Ó': 'O', 'Ỏ': 'O', 'Õ': 'O', 'Ọ': 'O',
                    'Ồ': 'O', 'Ố': 'O', 'Ổ': 'O', 'Ỗ': 'O', 'Ộ': 'O',
                    'Ù': 'U', 'Ú': 'U', 'Ủ': 'U', 'Ũ': 'U', 'Ụ': 'U',
                    'Ừ': 'U', 'Ứ': 'U', 'Ử': 'U', 'Ữ': 'U', 'Ự': 'U',
                    'Ỳ': 'Y', 'Ý': 'Y', 'Ỷ': 'Y', 'Ỹ': 'Y', 'Ỵ': 'Y',
                    'Â': 'A', 'Ă': 'A', 'Ê': 'E', 'Ô': 'O', 'Ư': 'U'
                }
                for vietnamese_char, replacement in vietnamese_map.items():
                    text = text.replace(vietnamese_char, replacement)
                return text
            except Exception as e:
                print(f"Text processing error: {e}")
                return str(text).encode('ascii', 'ignore').decode('ascii')
        
        # Tiêu đề
        c.setFont(font_bold, 16)
        title = f"Đề thi từ file: {original_filename}"
        try:
            title_width = c.stringWidth(title, font_bold, 16)
            c.drawString((width - title_width) / 2, height - 50, title)
        except:
            c.drawString(50, height - 50, title)
        
        y_position = height - 100
        c.setFont(font_name, 12)
        
        # Thêm các câu hỏi
        for item in quiz_data:
            if isinstance(item.get('question_number'), int):
                if y_position < 100:
                    c.showPage()
                    c.setFont(font_name, 12)
                    y_position = height - 50
                
                # Câu hỏi
                question_text = f"Câu {item['question_number']}: {item['question_text']}"
                
                # Chia text thành nhiều dòng nếu quá dài
                max_width = width - 100
                lines = []
                words = question_text.split(' ')
                current_line = ""
                
                for word in words:
                    test_line = current_line + word + " "
                    try:
                        line_width = c.stringWidth(test_line, font_name, 12)
                        if line_width < max_width:
                            current_line = test_line
                        else:
                            if current_line:
                                lines.append(current_line.strip())
                            current_line = word + " "
                    except:
                        if len(current_line) > 80:
                            if current_line:
                                lines.append(current_line.strip())
                            current_line = word + " "
                        else:
                            current_line = test_line
                
                if current_line:
                    lines.append(current_line.strip())
                
                # Vẽ câu hỏi
                for line in lines:
                    if y_position < 50:
                        c.showPage()
                        c.setFont(font_name, 12)
                        y_position = height - 50
                    c.drawString(50, y_position, line)
                    y_position -= 20
                
                # Thêm các lựa chọn
                for option in item.get('options', []):
                    if y_position < 50:
                        c.showPage()
                        c.setFont(font_name, 12)
                        y_position = height - 50
                    c.drawString(70, y_position, option)
                    y_position -= 15
                
                y_position -= 10  # Khoảng cách giữa các câu
        
        # Thêm bảng đáp án
        if any(item.get('correct_answer_char') for item in quiz_data):
            if y_position < 200:
                c.showPage()
                y_position = height - 50
            
            y_position -= 30
            c.setFont(font_bold, 14)
            c.drawString(50, y_position, "Bảng đáp án")
            y_position -= 30
            
            c.setFont(font_name, 10)
            quiz_with_answers = [item for item in quiz_data if isinstance(item.get('question_number'), int)]
            
            # Vẽ bảng theo hàng (4 câu mỗi hàng)
            for i in range(0, len(quiz_with_answers), 4):
                if y_position < 50:
                    c.showPage()
                    c.setFont(font_name, 10)
                    y_position = height - 50
                
                x_positions = [50, 120, 190, 260, 330, 400, 470, 540]
                
                for j in range(4):
                    if i + j < len(quiz_with_answers):
                        item = quiz_with_answers[i + j]
                        question_num = f"Câu {item['question_number']}"
                        answer_char = item.get('correct_answer_char', '?')
                        c.drawString(x_positions[j*2], y_position, question_num)
                        c.drawString(x_positions[j*2 + 1], y_position, answer_char)
                
                y_position -= 20
        
        c.save()
        
        # Đọc file và trả về
        with open(pdf_path, 'rb') as f:
            file_stream = io.BytesIO(f.read())
        
        os.unlink(pdf_path)
        
        file_stream.seek(0)
        base_name = original_filename.rsplit('.', 1)[0] if '.' in original_filename else original_filename
        download_filename = f"{base_name}_quiz.pdf"
        
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=download_filename,
            mimetype='application/pdf'
        )
        
    except Exception as e:
        print(f"Error creating PDF: {str(e)}")
        return jsonify({'error': f'Lỗi khi tạo file PDF: {str(e)}'}), 500
if __name__ == '__main__':
    app.run(debug=True)