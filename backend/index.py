# backend/api.py - FastAPI Backend cho UniAI

from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel
from typing import List, Optional
import os
import tempfile
import uuid
from dotenv import load_dotenv
from werkzeug.utils import secure_filename
import re
import random
import docx
import fitz  # PyMuPDF
import base64
import io
import html
import json
# Import từ agent_core hiện tại
from agent_core import load_documents, split_documents, create_vector_store, create_agent_executor,get_generation_llm,generate_essay_questions_logic 
from prompt_template import AGENT_SYSTEM_PROMPT,ESSAY_GENERATION_PROMPT_RAG, ESSAY_GENERATION_PROMPT_TOPIC
from podcast_generator import PodcastGenerator
from langchain_core.messages import HumanMessage, AIMessage

# 1. Load file .env trước
load_dotenv()

# 2. Xử lý Logic chọn API Key (Primary vs Backup)
def configure_google_api_key():
    primary_key = os.getenv("GOOGLE_API_KEY")
    backup_key = os.getenv("GOOGLE_API_KEY_BACKUP")

    # Kiểm tra xem key chính có tồn tại và không rỗng không
    if primary_key and primary_key.strip():
        print(f"[INFO] Đang sử dụng GOOGLE_API_KEY chính (***{primary_key[-4:]})")
        # LangChain tự nhận biến này, không cần set lại nếu nó đã đúng tên
    elif backup_key and backup_key.strip():
        print(f"[WARN] Key chính không tìm thấy hoặc rỗng. Chuyển sang GOOGLE_API_KEY_BACKUP (***{backup_key[-4:]})")
        # QUAN TRỌNG: Gán key dự phòng vào tên biến mà thư viện yêu cầu
        os.environ["GOOGLE_API_KEY"] = backup_key
    else:
        print("[CRITICAL] Không tìm thấy cả Primary Key lẫn Backup Key trong .env!")
        # Có thể raise lỗi nếu muốn app dừng ngay lập tức
        # raise ValueError("API Key is missing")

# Gọi hàm cấu hình ngay lập tức
configure_google_api_key()




app = FastAPI(title="UniAI Backend API", version="1.0.0")

# CORS middleware để cho phép React connect
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://localhost:5173"],  # React dev servers
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Cấu hình upload folder cho quiz
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx'}
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Lưu trữ sessions
sessions = {}

# --- HÀM HELPER CHO QUIZ ---
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def shuffle_questions(quiz_data):
    import random
    questions = [q for q in quiz_data if isinstance(q.get('question_number'), int)]
    other_items = [q for q in quiz_data if not isinstance(q.get('question_number'), int)]
    random.shuffle(questions)
    for i, question in enumerate(questions):
        question['question_number'] = i + 1
    return questions + other_items

def shuffle_answers_in_question(question):
    import random
    if not question.get('options') or not question['options']: return question
    clean_options = [re.sub(r'^[A-D][.\s]+', '', option.strip()) for option in question['options']]
    original_indices = list(range(len(clean_options)))
    random.shuffle(original_indices)
    shuffled_options = [clean_options[i] for i in original_indices]
    old_correct_index = question.get('correct_answer_index', 0)
    if old_correct_index < len(original_indices):
        new_correct_index = original_indices.index(old_correct_index)
        question['options'] = shuffled_options
        question['correct_answer_index'] = new_correct_index
        question['correct_answer_char'] = chr(65 + new_correct_index)
    return question

def shuffle_answers(quiz_data):
    return [shuffle_answers_in_question(q) if isinstance(q.get('question_number'), int) else q for q in quiz_data]

def snapshot_drawing(page, drawing_rect, dpi=150):
    try:
        clip = fitz.Rect(drawing_rect)
        pix = page.get_pixmap(clip=clip, dpi=dpi) 
        img_bytes = pix.tobytes("png")
        b64_string = base64.b64encode(img_bytes).decode('utf-8')
        return f'<img src="data:image/png;base64,{b64_string}" style="vertical-align: middle; max-height: 2.5em;" />'
    except Exception as e:
        print(f"Lỗi snapshot: {e}")
        return ""

def group_items_by_lines(items, threshold=5.0):
    lines = []
    if not items: return lines
    items.sort(key=lambda item: (item['bbox'].y0, item['bbox'].x0))
    current_line = [items[0]]
    last_y0 = items[0]['bbox'].y0
    for item in items[1:]:
        if abs(item['bbox'].y0 - last_y0) < threshold:
            current_line.append(item)
        else:
            current_line.sort(key=lambda x: x['bbox'].x0)
            lines.append(current_line)
            current_line = [item]
            last_y0 = item['bbox'].y0
    current_line.sort(key=lambda x: x['bbox'].x0)
    lines.append(current_line)
    return lines

def parse_horizontal_options(text):
    pattern = re.compile(r'([A-D][.:])(.*?)(?=\s*[A-D][.:]|\s*Câu\s*\d+|\Z)', re.IGNORECASE | re.DOTALL)
    matches = pattern.findall(text)
    return [f"{m[0].strip()} {m[1].strip()}" for m in matches]

# --- LOGIC XỬ LÝ DOCX ---
def extract_docx_data(file_path):
    try:
        doc = docx.Document(file_path)
        quiz_data = []

        current_quiz = None
        for p in doc.paragraphs:
            p_text = p.text.strip()
            if not p_text: continue

            match_question = re.match(r'^(Câu\s*(\d+)[:.]?)(.*)', p_text, re.IGNORECASE)
            if match_question:
                if current_quiz: quiz_data.append(current_quiz)
                current_quiz = {
                    "question_number": int(match_question.group(2)), 
                    "question_text": match_question.group(0), 
                    "options": [], 
                    "correct_answer_index": -1, 
                    "correct_answer_char": '', 
                    "images": []
                }
            elif current_quiz:
                if re.search(r'^[A-D][.:]', p_text):
                    current_quiz["options"].append(p_text.strip())
                else:
                    current_quiz["question_text"] += " " + p_text
                    
        if current_quiz: quiz_data.append(current_quiz)

        # Fallback strategy
        if not any(q['options'] for q in quiz_data):
            print("[DEBUG] DOCX: Fallback to single-line format strategy")
            quiz_data = []
            for p in doc.paragraphs:
                p_text = p.text.strip()
                if not p_text or not p_text.startswith("Câu"): continue
                
                question_blocks = re.split(r'(Câu\s*(\d+)[:.]?)', p_text, flags=re.IGNORECASE)
                if question_blocks[0] == '': question_blocks = question_blocks[1:]
                
                for i in range(0, len(question_blocks), 3):
                    if i + 2 >= len(question_blocks): continue
                    
                    question_header = question_blocks[i]
                    question_number = int(question_blocks[i+1])
                    question_body = question_blocks[i+2].strip()
                    
                    parts = re.split(r'(?=\s*[A-D][.:])', question_body, 1)
                    question_part = parts[0]
                    options_part = parts[1] if len(parts) > 1 else ""

                    options = parse_horizontal_options(options_part)
                    
                    full_question_text = (question_header + " " + question_part).strip()
                    quiz_data.append({
                        "question_number": question_number, 
                        "question_text": full_question_text,
                        "options": options, 
                        "correct_answer_index": -1, 
                        "correct_answer_char": '', 
                        "images": []
                    })

        # Tìm highlight
        for quiz in quiz_data:
            for p in doc.paragraphs:
                if f"Câu {quiz['question_number']}" in p.text:
                    for run in p.runs:
                        if run.font.highlight_color or (run.font.color and run.font.color.rgb):
                            highlighted_text = run.text.strip()
                            for i, option in enumerate(quiz["options"]):
                                if highlighted_text in option:
                                    quiz["correct_answer_char"] = chr(65 + i)
                                    quiz["correct_answer_index"] = i
                                    break
        
        # Dọn dẹp
        for quiz in quiz_data:
            if quiz["options"]:
                try:
                    if not quiz["correct_answer_char"] and quiz["options"]:
                        quiz["correct_answer_char"] = "A"
                        quiz["correct_answer_index"] = 0
                except (IndexError, TypeError): 
                    quiz["correct_answer_char"] = "A"
                    quiz["correct_answer_index"] = 0
                
                clean_options = []
                for option in quiz["options"]:
                    cleaned = re.sub(r'^[A-D][.\s]+', '', option.strip())
                    clean_options.append(cleaned)
                quiz["options"] = clean_options
            
            if quiz["correct_answer_char"]:
                quiz["correct_answer_index"] = ord(quiz["correct_answer_char"].upper()) - 65
        
        return quiz_data

    except Exception as e:
        print(f"Error reading docx: {str(e)}")
        return []

# --- LOGIC XỬ LÝ PDF ---
def extract_pdf_data(file_path):
    print(f"\n[DEBUG] Bắt đầu xử lý PDF: {file_path}")
    try:
        doc = fitz.open(file_path)
        quiz_data = []
        current_quiz = None
        current_state = "question" 
        all_lines = []
        
        # Thu thập tất cả dòng từ tất cả trang
        for page_num, page in enumerate(doc):
            print(f"[DEBUG] --- Gom dòng từ Trang {page_num + 1} / {len(doc)} ---")
            
            words = page.get_text("words")
            drawings = page.get_drawings()
            
            all_items = []
            for w in words:
                bbox = fitz.Rect(w[:4])
                all_items.append({"bbox": bbox, "type": "text", "content": f" {html.escape(w[4])} "})
            
            for d in drawings:
                bbox = fitz.Rect(d['rect'])
                if bbox.width < 5 or bbox.height < 5: continue
                img_html = snapshot_drawing(page, bbox)
                if img_html:
                    all_items.append({"bbox": bbox, "type": "image", "content": img_html})

            lines_on_page = group_items_by_lines(all_items, threshold=5.0)
            for line in lines_on_page:
                all_lines.append(line)
        
        print(f"[DEBUG] Đã gom tổng cộng {len(all_lines)} dòng từ tất cả các trang.")
        
        # Xử lý các dòng
        option_splitter_regex = re.compile(r'(\s*[A-D][.:])', re.IGNORECASE)
        
        for line_items in all_lines:
            if not line_items: continue
            
            line_html = "".join([item['content'] for item in line_items]).strip()
            line_text_only = "".join([item['content'] for item in line_items if item['type'] == 'text']).strip()
            line_text_only = re.sub(r'\s+', ' ', line_text_only).strip()
            
            if not line_text_only and not line_html: continue

            match_question = re.match(r'^(Câu\s*(\d+)[:.]?)', line_text_only, re.IGNORECASE)
            split_by_options_text = option_splitter_regex.split(line_text_only)
            
            if match_question:
                if current_quiz:
                    quiz_data.append(current_quiz)
                
                question_number = int(match_question.group(2))
                print(f"    [DEBUG] Tìm thấy: Câu {question_number}")
                
                question_part_html = line_html
                options_part_html_list = []
                
                split_by_options_html = option_splitter_regex.split(line_html)
                
                if len(split_by_options_html) > 1:
                    question_part_html = split_by_options_html[0]
                    for i in range(1, len(split_by_options_html), 2):
                        if i + 1 < len(split_by_options_html):
                            full_option_html = split_by_options_html[i] + split_by_options_html[i+1]
                            options_part_html_list.append(full_option_html)
                
                current_quiz = {
                    "question_number": question_number, 
                    "question_text": question_part_html,
                    "options": options_part_html_list,
                    "correct_answer_index": -1, 
                    "correct_answer_char": '', 
                    "images": []
                }
                
                current_state = "option" if options_part_html_list else "question"

            elif current_quiz: 
                if len(split_by_options_text) > 1 and not re.match(r'^[A-D][.:]', split_by_options_text[0].strip()):
                    split_by_options_html = option_splitter_regex.split(line_html)
                    
                    first_part = split_by_options_html[0].strip()
                    if first_part:
                        if current_state == "option" and current_quiz["options"]:
                            current_quiz["options"][-1] += " " + first_part
                        else:
                            current_quiz["question_text"] += " " + first_part
                    
                    for i in range(1, len(split_by_options_html), 2):
                        if i + 1 < len(split_by_options_html):
                            full_option_html = split_by_options_html[i] + split_by_options_html[i+1]
                            current_quiz["options"].append(full_option_html)
                    
                    current_state = "option"
                
                elif re.match(r'^[A-D][.:]', line_text_only):
                    current_quiz["options"].append(line_html)
                    current_state = "option"

                else:
                    if current_state == "option" and current_quiz["options"]:
                        current_quiz["options"][-1] += " " + line_html
                    else:
                        current_quiz["question_text"] += " " + line_html
                
        if current_quiz:
            quiz_data.append(current_quiz)
        
        doc.close()

        # Dọn dẹp
        for quiz in quiz_data:
            if quiz["options"]:
                if not quiz["correct_answer_char"]:
                    quiz["correct_answer_char"] = "A"
                    quiz["correct_answer_index"] = 0
                
                clean_options = []
                for option_html in quiz["options"]:
                    cleaned_html = re.sub(
                        r'^(<[^>]+>)*\s*[A-D][.:]\s*', 
                        r'\1', 
                        option_html.strip(), 
                        count=1, 
                        flags=re.IGNORECASE
                    )
                    clean_options.append(cleaned_html)
                quiz["options"] = clean_options
        
        print(f"[DEBUG] Hoàn tất, tổng cộng {len(quiz_data)} câu.")
        return quiz_data

    except Exception as e:
        print(f"[ERROR] Lỗi nghiêm trọng khi xử lý PDF: {str(e)}")
        import traceback
        traceback.print_exc() 
        return []

def download_quiz_docx(quiz_data, original_filename):
    """Tạo file Word chứa đề thi"""
    try:
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = docx.Document()
        
        title = doc.add_heading(f'Đề thi từ file: {original_filename}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for item in quiz_data:
            if isinstance(item.get('question_number'), int):
                question_para = doc.add_paragraph()
                question_para.add_run(f"Câu {item['question_number']}: ").bold = True
                question_para.add_run(item['question_text'])
                
                for i, option in enumerate(item.get('options', [])):
                    option_para = doc.add_paragraph(f"{chr(65+i)}. {option}")
                
                doc.add_paragraph()
        
        if any(item.get('correct_answer_char') for item in quiz_data):
            doc.add_heading('Bảng đáp án', level=1)
            
            table = doc.add_table(rows=1, cols=8)
            table.style = 'Table Grid'
            
            header_cells = table.rows[0].cells
            headers = ['Câu', 'Đáp án', 'Câu', 'Đáp án', 'Câu', 'Đáp án', 'Câu', 'Đáp án']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            quiz_with_answers = [item for item in quiz_data if isinstance(item.get('question_number'), int)]
            for i in range(0, len(quiz_with_answers), 4):
                row_cells = table.add_row().cells
                for j in range(4):
                    if i + j < len(quiz_with_answers):
                        item = quiz_with_answers[i + j]
                        row_cells[j*2].text = str(item['question_number'])
                        row_cells[j*2+1].text = item.get('correct_answer_char', '?')
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        base_name = original_filename.rsplit('.', 1)[0] if '.' in original_filename else original_filename
        download_filename = f"{base_name}_quiz.docx"
        
        return file_stream, download_filename
        
    except Exception as e:
        print(f"Error creating DOCX: {str(e)}")
        raise HTTPException(status_code=500, detail=f'Lỗi khi tạo file Word: {str(e)}')

def download_quiz_pdf(quiz_data, original_filename):
    """Tạo file PDF chứa đề thi"""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            pdf_path = tmp_file.name
        
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            from reportlab.lib.units import inch
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
        except ImportError:
            raise HTTPException(status_code=500, detail='Thư viện reportlab chưa được cài đặt')
        
        c = canvas.Canvas(pdf_path, pagesize=A4)
        width, height = A4
        
        font_name = 'Helvetica'
        font_bold = 'Helvetica-Bold'
        
        def safe_text(text):
            if not text: return ""
            try:
                return str(text).encode('utf-8').decode('utf-8')
            except:
                return ''.join(char for char in str(text) if ord(char) < 128)
        
        c.setFont(font_bold, 16)
        title = f"Đề thi từ file: {original_filename}"
        try:
            title_width = c.stringWidth(title, font_bold, 16)
            c.drawString((width - title_width) / 2, height - 50, safe_text(title))
        except:
            c.drawString(50, height - 50, safe_text(title))
        
        y_position = height - 100
        c.setFont(font_name, 12)
        
        for item in quiz_data:
            if isinstance(item.get('question_number'), int):
                if y_position < 150:
                    c.showPage()
                    y_position = height - 50
                
                question_text = safe_text(f"Câu {item['question_number']}: {item['question_text']}")
                c.drawString(50, y_position, question_text[:100] + "..." if len(question_text) > 100 else question_text)
                y_position -= 20
                
                for i, option in enumerate(item.get('options', [])):
                    if y_position < 50:
                        c.showPage()
                        y_position = height - 50
                    
                    option_text = safe_text(f"{chr(65+i)}. {option}")
                    c.drawString(70, y_position, option_text[:80] + "..." if len(option_text) > 80 else option_text)
                    y_position -= 15
                
                y_position -= 10
        
        c.save()
        
        with open(pdf_path, 'rb') as f:
            file_stream = io.BytesIO(f.read())
        
        os.unlink(pdf_path)
        
        file_stream.seek(0)
        base_name = original_filename.rsplit('.', 1)[0] if '.' in original_filename else original_filename
        download_filename = f"{base_name}_quiz.pdf"
        
        return file_stream, download_filename
        
    except Exception as e:
        print(f"Error creating PDF: {str(e)}")
        raise HTTPException(status_code=500, detail=f'Lỗi khi tạo file PDF: {str(e)}')

class ChatMessage(BaseModel):
    content: str
    role: str  # "user" hoặc "assistant"

class ChatRequest(BaseModel):
    message: str
    session_id: str

class ChatResponse(BaseModel):
    response: str
    is_quiz: bool = False

class DocumentUploadResponse(BaseModel):
    message: str
    session_id: str
    success: bool



class PodcastRequest(BaseModel):
    session_id: str

class PodcastResponse(BaseModel):
    success: bool
    dialogue: Optional[str] = None
    audio_url: Optional[str] = None
    message: str

class QuizExtractionResponse(BaseModel):
    success: bool
    filename: str
    quiz_data: List[dict]
    total_questions: int

class QuizShuffleRequest(BaseModel):
    quiz_data: List[dict]
    shuffle_questions: bool = False
    shuffle_answers: bool = False

class QuizShuffleResponse(BaseModel):
    success: bool
    quiz_data: List[dict]
    total_questions: int

class QuizDownloadRequest(BaseModel):
    quiz_data: List[dict]
    filename: str
    format: str  # 'pdf' or 'docx'


class EssayQuestionItem(BaseModel):
    question_number: int
    question_text: str
    suggested_answer: str
class EssayGenerationRequest(BaseModel):
    num_questions: int
    session_id: Optional[str] = None
    topic: Optional[str] = None

class EssayGenerationResponse(BaseModel):
    success: bool
    questions: List[EssayQuestionItem] = []
    message: Optional[str] = None
    
    
    
@app.get("/")
async def root():
    return {"message": "UniAI Backend API is running!"}

@app.post("/upload-documents", response_model=DocumentUploadResponse)
async def upload_documents(files: List[UploadFile] = File(...)):
    """Upload và xử lý tài liệu PDF"""
    try:
        session_id = str(uuid.uuid4())
        
        # Lưu files tạm thời
        temp_files = []
        uploaded_files = []
        
        for file in files:
            if not file.filename.endswith('.pdf'):
                raise HTTPException(status_code=400, detail=f"File {file.filename} không phải PDF")
            
            # Tạo file tạm
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            content = await file.read()
            temp_file.write(content)
            temp_file.close()
            
            temp_files.append(temp_file.name)
            
            # Tạo object giống như streamlit file uploader
            class MockFile:
                def __init__(self, name, path):
                    self.name = name
                    self.path = path
                
                def getbuffer(self):
                    with open(self.path, 'rb') as f:
                        return f.read()
            
            uploaded_files.append(MockFile(file.filename, temp_file.name))
        
        # Xử lý documents
        docs = load_documents(uploaded_files)
        chunks = split_documents(docs)
        vector_store = create_vector_store(chunks)
        
        # Tạo agent executor
        agent_executor = create_agent_executor(
            vector_store, 
            AGENT_SYSTEM_PROMPT, 
            text_chunks=chunks
        )
        
        # Lưu session
        sessions[session_id] = {
            'agent_executor': agent_executor,
            'text_chunks': chunks,  # Thêm text_chunks để dùng cho podcast
            'chat_history': [],
            'processed_files': [f.filename for f in files]
        }
        
        # Cleanup temp files
        for temp_file in temp_files:
            os.unlink(temp_file)
        
        return DocumentUploadResponse(
            message=f"Đã xử lý thành công {len(files)} tài liệu",
            session_id=session_id,
            success=True
        )
        
    except Exception as e:
        # Cleanup temp files nếu có lỗi
        for temp_file in temp_files:
            if os.path.exists(temp_file):
                os.unlink(temp_file)
        
        raise HTTPException(status_code=500, detail=f"Lỗi xử lý tài liệu: {str(e)}")

@app.post("/chat", response_model=ChatResponse)
async def chat(request: ChatRequest):
    """Xử lý chat với UniAI"""
    try:
        print(f"[DEBUG] Received chat request: {request.message[:100]}...")
        print(f"[DEBUG] Session ID: {request.session_id}")
        print(f"[DEBUG] Available sessions: {list(sessions.keys())}")
        
        if request.session_id not in sessions:
            print(f"[ERROR] Session not found: {request.session_id}")
            print(f"[ERROR] Available sessions: {list(sessions.keys())}")
            raise HTTPException(status_code=404, detail="Session không tồn tại. Vui lòng upload lại tài liệu.")
        
        session = sessions[request.session_id]
        agent_executor = session['agent_executor']
        chat_history = session['chat_history']
        
        print(f"[DEBUG] Chat history length: {len(chat_history)}")
        print(f"[DEBUG] Agent executor type: {type(agent_executor)}")
        
        # Thêm message của user vào history
        user_message = HumanMessage(content=request.message)
        chat_history.append(user_message)
        
        print(f"[DEBUG] Calling agent with input: {request.message[:100]}...")
        
        # Gọi agent với better error handling
        try:
            response = agent_executor.invoke({
                "input": request.message,
                "chat_history": chat_history
            })
            print(f"[DEBUG] Agent execution successful")
        except Exception as agent_error:
            print(f"[ERROR] Agent execution failed: {agent_error}")
            # Fallback response
            answer = f"Xin lỗi, có lỗi xảy ra khi xử lý câu hỏi của bạn. Lỗi: {str(agent_error)}. Vui lòng thử lại hoặc upload lại tài liệu."
            return ChatResponse(response=answer, is_quiz=False)
        
        answer = response['output']
        print(f"[DEBUG] Agent response length: {len(answer)}")
        print(f"[DEBUG] Agent response preview: {answer[:200]}...")
        
        # Thêm response của AI vào history
        ai_message = AIMessage(content=answer)
        chat_history.append(ai_message)
        
        # Cập nhật session
        sessions[request.session_id]['chat_history'] = chat_history
        
        # Kiểm tra nếu là quiz bằng cách phân tích nội dung
        is_quiz = False
        try:
            # Chỉ parse JSON nếu response có dấu hiệu là JSON
            import json
            answer_stripped = answer.strip()
            if answer_stripped.startswith('{') and answer_stripped.endswith('}'):
                parsed = json.loads(answer)
                if isinstance(parsed, dict) and 'questions' in parsed and 'quiz_title' in parsed:
                    is_quiz = True
                    print(f"[DEBUG] Detected valid quiz JSON with {len(parsed.get('questions', []))} questions")
            else:
                # Không phải JSON, kiểm tra bằng từ khóa trong request
                is_quiz = any(keyword in request.message.lower() 
                             for keyword in ["quiz", "trắc nghiệm", "câu hỏi", "test", "kiểm tra"])
                
                # Hoặc kiểm tra trong response có chứa format quiz
                if not is_quiz:
                    quiz_indicators = ['"questions":', '"quiz_title":', '"correct_answer":']
                    is_quiz = any(indicator in answer for indicator in quiz_indicators)
                    
                if is_quiz:
                    print(f"[DEBUG] Detected quiz by keywords/indicators")
                else:
                    print(f"[DEBUG] Normal response - not a quiz")
                    
        except Exception as json_error:
            # Fallback: chỉ dựa vào từ khóa
            print(f"[DEBUG] JSON parse error (expected for normal responses): {json_error}")
            is_quiz = any(keyword in request.message.lower() 
                         for keyword in ["quiz", "trắc nghiệm", "câu hỏi", "test", "kiểm tra"])
            print(f"[DEBUG] Fallback quiz detection by keywords: {is_quiz}")
        
        print(f"[DEBUG] Final is_quiz determination: {is_quiz}")
        
        return ChatResponse(
            response=answer,
            is_quiz=is_quiz
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Lỗi xử lý chat: {str(e)}")

@app.get("/session/{session_id}/info")
async def get_session_info(session_id: str):
    """Lấy thông tin session"""
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session không tồn tại")
    
    session = sessions[session_id]
    return {
        "session_id": session_id,
        "processed_files": session.get('processed_files', []),
        "chat_count": len(session.get('chat_history', [])),
        "agent_status": "active" if session.get('agent_executor') else "inactive",
        "session_age": "unknown"  # Could add timestamp tracking
    }

@app.get("/sessions")
async def list_sessions():
    """List all active sessions - for debugging"""
    return {
        "active_sessions": list(sessions.keys()),
        "session_count": len(sessions)
    }

@app.get("/session/{session_id}/history")
async def get_chat_history(session_id: str):
    """Lấy lịch sử chat"""
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session không tồn tại")
    
    chat_history = sessions[session_id].get('chat_history', [])
    
    # Convert LangChain messages to simple dict
    formatted_history = []
    for msg in chat_history:
        if isinstance(msg, HumanMessage):
            formatted_history.append({"role": "user", "content": msg.content})
        elif isinstance(msg, AIMessage):
            formatted_history.append({"role": "assistant", "content": msg.content})
    
    return {"chat_history": formatted_history}


@app.post("/generate-podcast", response_model=PodcastResponse)
async def generate_podcast(request: PodcastRequest):
    """Tạo podcast từ nội dung PDF đã upload"""
    try:
        print(f"[DEBUG] Generating podcast for session: {request.session_id}")
        
        if request.session_id not in sessions:
            print(f"[ERROR] Session {request.session_id} not found")
            raise HTTPException(status_code=404, detail="Session không tồn tại")
        
        session = sessions[request.session_id]
        print(f"[DEBUG] Session keys: {list(session.keys())}")
        
        # Lấy nội dung văn bản từ session
        text_chunks = session.get('text_chunks', [])
        print(f"[DEBUG] Found {len(text_chunks)} text chunks")
        
        if not text_chunks:
            raise HTTPException(status_code=400, detail="Không có nội dung để tạo podcast")
        
        # Gộp tất cả chunks thành một văn bản
        pdf_content = "\n\n".join([chunk.page_content for chunk in text_chunks])
        print(f"[DEBUG] PDF content length: {len(pdf_content)} characters")
        
        # Giới hạn độ dài nội dung (để tránh token limit)
        max_chars = 10000  # Giới hạn 10k ký tự
        if len(pdf_content) > max_chars:
            pdf_content = pdf_content[:max_chars] + "..."
            print(f"[DEBUG] Content truncated to {max_chars} characters")
        
        # Tạo podcast
        print("[DEBUG] Creating PodcastGenerator...")
        podcast_gen = PodcastGenerator()
        print("[DEBUG] Generating podcast...")
        result = podcast_gen.generate_podcast(pdf_content)
        
        print(f"[DEBUG] Podcast generation result: {result.get('success', False)}")
        
        if result['success']:
            return PodcastResponse(
                success=True,
                dialogue=result['dialogue'],
                audio_url=f"/audio/{os.path.basename(result['audio_path'])}",  # URL để download audio
                message=result['message']
            )
        else:
            return PodcastResponse(
                success=False,
                message=result['message']
            )
            
    except HTTPException:
        raise
    except Exception as e:
        print(f"[ERROR] Error generating podcast: {str(e)}")
        import traceback
        traceback.print_exc()
        return PodcastResponse(
            success=False,
            message=f"Lỗi khi tạo podcast: {str(e)}"
        )

@app.get("/audio/{filename}")
async def get_audio(filename: str):
    """Serve audio files"""
    audio_path = os.path.join(tempfile.gettempdir(), filename)
    if os.path.exists(audio_path):
        return FileResponse(
            audio_path,
            media_type="audio/mpeg",
            headers={"Content-Disposition": f"inline; filename={filename}"}
        )
    else:
        raise HTTPException(status_code=404, detail="Audio file không tìm thấy")

# --- QUIZ API ENDPOINTS ---



# --- MỚI: ENDPOINT TẠO CÂU HỎI TỰ LUẬN ---
@app.post("/api/generate-essay-questions", response_model=EssayGenerationResponse)
async def generate_essay_questions_api(request: EssayGenerationRequest):
    """
    API endpoint để tạo câu hỏi tự luận.
    - Nếu có `session_id`: Tạo câu hỏi từ file PDF đã upload (RAG).
    - Nếu có `topic`: Tạo câu hỏi từ chủ đề (General Knowledge).
    """
    try:
        print(f"[DEBUG] Nhận yêu cầu tạo {request.num_questions} câu hỏi tự luận.")

        # --- Kiểm tra đầu vào ---
        if not request.session_id and not request.topic:
            raise HTTPException(status_code=400, detail="Vui lòng cung cấp 'session_id' (cho file) hoặc 'topic' (cho chủ đề).")
        
        if request.session_id and request.topic:
            raise HTTPException(status_code=400, detail="Chỉ cung cấp 'session_id' hoặc 'topic', không cung cấp cả hai.")

        if request.num_questions <= 0:
            raise HTTPException(status_code=400, detail="Số lượng câu hỏi phải lớn hơn 0.")

        # Khởi tạo LLM
        llm = get_generation_llm()
        
        generated_questions = []

        # --- TRƯỜNG HỢP 1: TẠO TỪ FILE (RAG) ---
        if request.session_id:
            print(f"[DEBUG] Tạo tự luận từ session_id: {request.session_id}")
            if request.session_id not in sessions:
                raise HTTPException(status_code=404, detail="Session không tồn tại. Vui lòng upload lại tài liệu.")
            
            session = sessions[request.session_id]
            text_chunks = session.get('text_chunks', [])
            
            if not text_chunks:
                raise HTTPException(status_code=400, detail="Không có nội dung tài liệu trong session này để tạo câu hỏi.")
            
            # Gộp tất cả text chunks thành một context lớn
            full_context = "\n\n---\n\n".join([chunk.page_content for chunk in text_chunks])
            print(f"[DEBUG] Tổng độ dài context: {len(full_context)} ký tự")

            # Gọi hàm logic (async)
            generated_questions = await generate_essay_questions_logic(
                llm=llm,
                prompt_template_str=ESSAY_GENERATION_PROMPT_RAG,
                num_questions=request.num_questions,
                context=full_context
            )

        # --- TRƯỜNG HỢP 2: TẠO TỪ CHỦ ĐỀ ---
        elif request.topic:
            print(f"[DEBUG] Tạo tự luận từ chủ đề: {request.topic}")
            
            # Gọi hàm logic (async)
            generated_questions = await generate_essay_questions_logic(
                llm=llm,
                prompt_template_str=ESSAY_GENERATION_PROMPT_TOPIC,
                num_questions=request.num_questions,
                topic=request.topic
            )

        # --- Trả về kết quả ---
        if generated_questions:
            return EssayGenerationResponse(
                success=True,
                questions=[EssayQuestionItem(**q) for q in generated_questions]
            )
        else:
            # Trường hợp này hiếm khi xảy ra nếu logic trên chạy đúng
            raise HTTPException(status_code=500, detail="Không thể tạo câu hỏi từ thông tin được cung cấp.")

    except ValueError as ve:
        # Lỗi logic, ví dụ: parse JSON thất bại
        print(f"[ERROR] ValueError in generate_essay_questions_api: {str(ve)}")
        return EssayGenerationResponse(success=False, message=str(ve))
    except HTTPException as he:
        # Re-raise HTTPException (lỗi 400, 404, v.v.)
        raise he
    except Exception as e:
        # Các lỗi 500 khác
        print(f"[ERROR] Lỗi nghiêm trọng trong /api/generate-essay-questions: {str(e)}")
        import traceback
        traceback.print_exc()
        return EssayGenerationResponse(
            success=False, 
            message=f"Lỗi máy chủ nội bộ: {str(e)}"
        )



@app.get("/api/health")
async def quiz_health_check():
    """Health check endpoint cho quiz API"""
    return {"status": "ok", "message": "Quiz API is running"}

@app.post("/api/extract-quiz", response_model=QuizExtractionResponse)
async def extract_quiz_api(file: UploadFile = File(...)):
    """API endpoint để trích xuất câu hỏi từ file PDF/DOCX"""
    try:
        if not file.filename:
            raise HTTPException(status_code=400, detail='Không có file được chọn')
            
        if not allowed_file(file.filename):
            raise HTTPException(status_code=400, detail='Định dạng file không được hỗ trợ. Chỉ chấp nhận .pdf và .docx')
        
        filename = secure_filename(file.filename)
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        
        # Lưu file
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        print(f"\n[DEBUG] Đã lưu file: {file_path}")
        
        try:
            quiz_data = []
            file_ext = filename.rsplit('.', 1)[1].lower()
            print(f"[DEBUG] Loại file: {file_ext}")
            
            if file_ext == 'docx':
                print("[DEBUG] Đang gọi extract_docx_data...")
                quiz_data = extract_docx_data(file_path)
            elif file_ext == 'pdf':
                print("[DEBUG] Đang gọi extract_pdf_data...")
                quiz_data = extract_pdf_data(file_path)
            
            print(f"[DEBUG] Trích xuất xong, tìm thấy {len(quiz_data)} câu hỏi.")
            
            os.remove(file_path)
            print(f"[DEBUG] Đã xóa file tạm: {file_path}")
            
            return QuizExtractionResponse(
                success=True,
                filename=filename,
                quiz_data=quiz_data,
                total_questions=len([q for q in quiz_data if isinstance(q.get('question_number'), int)])
            )
            
        except Exception as e:
            if os.path.exists(file_path):
                os.remove(file_path)
            raise e
            
    except Exception as e:
        print(f"Error in extract_quiz_api: {str(e)}")
        raise HTTPException(status_code=500, detail=f'Lỗi khi xử lý file: {str(e)}')

@app.post("/api/shuffle-quiz", response_model=QuizShuffleResponse)
async def shuffle_quiz_api(request: QuizShuffleRequest):
    """API endpoint để áp dụng đảo câu hỏi và đáp án"""
    try:
        quiz_data = request.quiz_data
        shuffle_questions_enabled = request.shuffle_questions
        shuffle_answers_enabled = request.shuffle_answers
        
        if shuffle_answers_enabled:
            quiz_data = shuffle_answers(quiz_data)
        
        if shuffle_questions_enabled:
            quiz_data = shuffle_questions(quiz_data)
        
        return QuizShuffleResponse(
            success=True,
            quiz_data=quiz_data,
            total_questions=len([q for q in quiz_data if isinstance(q.get('question_number'), int)])
        )
        
    except Exception as e:
        print(f"Error in shuffle_quiz_api: {str(e)}")
        raise HTTPException(status_code=500, detail=f'Lỗi khi đảo đề: {str(e)}')

@app.post("/api/download-quiz")
async def download_quiz_api(request: QuizDownloadRequest):
    """API endpoint để tải về đề thi"""
    try:
        quiz_data = request.quiz_data
        filename = request.filename
        format_type = request.format
        
        if format_type == 'pdf':
            file_stream, download_filename = download_quiz_pdf(quiz_data, filename)
            file_stream.seek(0)
            return StreamingResponse(
                io.BytesIO(file_stream.read()),
                media_type='application/pdf',
                headers={"Content-Disposition": f"attachment; filename={download_filename}"}
            )
        elif format_type == 'docx':
            file_stream, download_filename = download_quiz_docx(quiz_data, filename)
            file_stream.seek(0)
            return StreamingResponse(
                io.BytesIO(file_stream.read()),
                media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                headers={"Content-Disposition": f"attachment; filename={download_filename}"}
            )
        else:
            raise HTTPException(status_code=400, detail='Định dạng không được hỗ trợ')
            
    except Exception as e:
        print(f"Error in download_quiz_api: {str(e)}")
        raise HTTPException(status_code=500, detail=f'Lỗi khi tạo file: {str(e)}')

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("api:app", host="0.0.0.0", port=8000, reload=True)
